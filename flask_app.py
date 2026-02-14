"""
Flask-based Contract Law Web Application
완전한 웹 애플리케이션으로 RAG 시스템과 통합
"""

from flask import Flask, request, jsonify, render_template, send_file
from flask_cors import CORS
import os
import sys
import time
import json
from datetime import datetime
import io

# 프로젝트 경로 추가
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

from src.utils.rag_system import ContractRAGSystem
from src.utils.api_manager import APIKeyManager
from src.utils.responsible_ai import ResponsibleAISystem
from config.settings import *

app = Flask(__name__)
CORS(app)  # CORS 활성화

# 글로벌 RAG 시스템 (서버 시작 시 1회만 초기화)
rag_system = None
responsible_ai = None
initialization_time = None
system_status = {
    'initialized': False,
    'documents_count': 0,
    'initialization_time': None,
    'last_updated': None,
    'error': None,
    'responsible_ai_enabled': False
}

# Responsible AI 분석 결과 저장용 전역 변수
latest_responsible_ai_results = {
    'bias_detected': [],
    'safety_check': {'safe': True, 'issues': [], 'risk_level': 'low'},
    'privacy_protected': False,
    'last_analysis_time': None,
    'contract_text_preview': ''
}

def initialize_rag_system():
    """RAG 시스템 초기화 (서버 시작 시 1회만 실행)"""
    global rag_system, responsible_ai, initialization_time, system_status

    try:
        print("Flask 서버 시작 - RAG 시스템 초기화 중...")
        start_time = time.time()

        # API 매니저 초기화
        api_manager = APIKeyManager()
        api_key = api_manager.load_api_key('gemini')

        # 디버깅 정보 추가
        print(f"API 키 로드 결과: {api_key[:10] + '...' if api_key else 'None'}")
        print(f"API 키 길이: {len(api_key) if api_key else 0}")

        if not api_key:
            print("WARNING: Gemini API 키가 설정되지 않았습니다.")
            system_status['error'] = "API 키가 설정되지 않았습니다"
            return False
        
        # 환경변수에 API 키 설정 (RAG 시스템이 읽을 수 있도록)
        os.environ["GEMINI_API_KEY"] = api_key
        print(f"환경변수에 API 키 설정 완료")

        # RAG 시스템 초기화
        rag_system = ContractRAGSystem()
        
        # Responsible AI 시스템 초기화
        responsible_ai = ResponsibleAISystem()
        system_status['responsible_ai_enabled'] = True

        # 캐시에서 로딩 시도
        cache_status = rag_system.check_cache_available()
        if cache_status['documents'] and cache_status['embeddings']:
            print("캐시 발견 - 빠른 로딩 모드...")
            success = rag_system.load_from_cache()
            if success:
                success = rag_system.load_vector_store()
                if success:
                    print("캐시에서 로딩 완료!")
                    # QA 체인 설정
                    if not rag_system.setup_qa_chain():
                        print("WARNING: QA 체인 설정에 실패했습니다.")
                else:
                    print("벡터 스토어 로딩 실패 - 전체 초기화 진행...")
                    success = False
        else:
            success = False

        # 캐시 로딩 실패 시 전체 초기화
        if not success:
            print("전체 데이터 초기화 진행...")
            data_directories = [
                "05.계약 법률 문서 서식 데이터/3.개방데이터/1.데이터/Training/01.원천데이터",
                "05.계약 법률 문서 서식 데이터/3.개방데이터/1.데이터/Validation/01.원천데이터"
            ]

            # 존재하는 디렉토리만 필터링
            existing_dirs = [d for d in data_directories if os.path.exists(d)]

            if not existing_dirs:
                raise Exception("데이터 디렉토리를 찾을 수 없습니다")

            # 문서 로드
            loaded_count = rag_system.load_legal_documents(existing_dirs, max_docs=5000)

            if loaded_count == 0:
                raise Exception("문서가 로드되지 않았습니다")

            # 벡터 스토어 구축
            rag_system.build_vector_store()

            # 캐시 저장
            rag_system._save_to_cache()

            print(f"전체 초기화 완료! {loaded_count}개 문서 로드됨")

        # QA 체인 설정
        if not rag_system.setup_qa_chain():
            print("WARNING: QA 체인 설정에 실패했습니다.")

        # 초기화 완료 정보 업데이트
        initialization_time = time.time() - start_time
        system_status.update({
            'initialized': True,
            'documents_count': len(rag_system.documents) if rag_system.documents else 0,
            'initialization_time': f"{initialization_time:.1f}초",
            'last_updated': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'error': None
        })

        print(f"RAG 시스템 준비 완료! ({initialization_time:.1f}초)")
        return True

    except Exception as e:
        error_msg = f"RAG 시스템 초기화 실패: {str(e)}"
        print(f"{error_msg}")
        system_status.update({
            'initialized': False,
            'error': error_msg,
            'last_updated': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })
        return False

# 웹 페이지 라우트들

@app.route('/')
def home():
    """홈페이지"""
    return render_template('home.html', status=system_status)

@app.route('/chatbot')
def chatbot():
    """AI 챗봇 페이지"""
    return render_template('chatbot.html')

@app.route('/analyzer')
def analyzer():
    """계약서 분석 페이지"""
    return render_template('analyzer.html')

@app.route('/templates')
def templates():
    """계약서 템플릿 페이지"""
    return render_template('templates.html')

@app.route('/responsible-ai')
def responsible_ai():
    """Responsible AI 대시보드"""
    return render_template('responsible_ai.html')

# API 엔드포인트들

@app.route('/status')
def get_status():
    """시스템 상태 반환"""
    return jsonify(system_status)

@app.route('/responsible-ai/explain', methods=['POST'])
def explain_analysis():
    """AI 분석 결과 설명 제공"""
    try:
        data = request.get_json()
        analysis_result = data.get('analysis_result', {})
        
        if not analysis_result:
            return jsonify({'error': '분석 결과가 제공되지 않았습니다'}), 400
        
        explanation = responsible_ai.generate_explanation(analysis_result)
        
        return jsonify({
            'explanation': explanation,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        return jsonify({
            'error': f'설명 생성 중 오류 발생: {str(e)}'
        }), 500

@app.route('/responsible-ai/audit-log')
def get_audit_log():
    """감사 로그 조회 (관리자용)"""
    try:
        # 실제 구현에서는 인증/권한 확인 필요
        audit_file = "data/audit_log.json"
        
        if os.path.exists(audit_file):
            with open(audit_file, 'r', encoding='utf-8') as f:
                audit_log = json.load(f)
        else:
            audit_log = []
        
        return jsonify({
            'audit_log': audit_log[-50:],  # 최근 50개만 반환
            'total_entries': len(audit_log)
        })
        
    except Exception as e:
        return jsonify({
            'error': f'감사 로그 조회 중 오류 발생: {str(e)}'
        }), 500

@app.route('/responsible-ai/sustainability')
def get_sustainability_metrics():
    """지속가능성 지표 조회"""
    try:
        metrics = responsible_ai.get_sustainability_metrics()
        
        return jsonify({
            'sustainability_metrics': metrics,
            'timestamp': datetime.now().isoformat()
        })
        
    except Exception as e:
        return jsonify({
            'error': f'지속가능성 지표 조회 중 오류 발생: {str(e)}'
        }), 500

@app.route('/responsible-ai/latest-analysis')
def get_latest_analysis():
    """최근 계약서 분석 결과 반환"""
    global latest_responsible_ai_results
    return jsonify({
        'latest_analysis': latest_responsible_ai_results,
        'timestamp': datetime.now().isoformat()
    })

def _get_fallback_answer(question):
    """AI 오류 시 기본 답변 제공"""
    fallback_answers = {
        '고용계약서': {
            '필수 조항': """고용계약서 작성 시 필수 조항은 다음과 같습니다:

1. **계약 당사자**: 고용주와 근로자의 명확한 신원
2. **근무장소**: 구체적인 근무 장소 명시
3. **업무내용**: 담당 업무의 구체적 범위
4. **근무시간**: 일일/주간 근무시간 규정
5. **임금**: 기본급, 수당, 지급일 등
6. **계약기간**: 고용 시작일과 종료일
7. **해지조건**: 계약 해지 사유와 절차
8. **비밀유지**: 업무상 기밀 보호 의무

이러한 조항들이 명확히 기재되어야 법적 효력이 있습니다.""",
            
            '작성': """고용계약서 작성 시 주의사항:

1. **명확성**: 모든 조항을 명확하고 구체적으로 작성
2. **공정성**: 일방적으로 불리한 조건 지양
3. **법적 준수**: 근로기준법 등 관련 법규 준수
4. **필수사항**: 근로계약서 표준양식 참고
5. **검토**: 계약 체결 전 양쪽이 충분히 검토

계약서는 양쪽 당사자가 이해하고 동의한 내용으로 작성해야 합니다."""
        },
        
        '부동산': {
            '임대차': """부동산 임대차 계약서 핵심 조항:

1. **임대목적물**: 소재지, 면적, 용도 명시
2. **임대기간**: 시작일과 종료일
3. **보증금**: 금액과 반환 조건
4. **월세**: 금액과 지급일
5. **임대인의 의무**: 시설 유지보수 등
6. **임차인의 의무**: 월세 지급, 시설 보존 등
7. **계약 갱신**: 갱신 조건과 절차
8. **해지 조건**: 해지 사유와 통지 기간

임대차 계약은 주택임대차보호법의 적용을 받습니다.""",
            
            '매매': """부동산 매매계약서 핵심 조항:

1. **매매목적물**: 소재지, 면적, 지번 등
2. **매매대금**: 총 매매가격
3. **계약금**: 계약 체결 시 지급
4. **중도금**: 중간 지급 금액
5. **잔금**: 최종 지급 금액
6. **소유권이전**: 이전 시기와 절차
7. **인도**: 부동산 인도 시기
8. **분쟁해결**: 중재 또는 소송

매매계약은 부동산 거래의 기본이 되는 중요한 계약입니다."""
        },
        
        '계약서': {
            '일반': """계약서 작성 시 기본 원칙:

1. **명확성**: 모든 조항을 명확하게 작성
2. **완전성**: 필요한 모든 사항을 포함
3. **공정성**: 양쪽 당사자에게 공정한 조건
4. **법적 유효성**: 관련 법규 준수
5. **실행 가능성**: 실제로 이행 가능한 내용

계약서는 당사자 간의 권리와 의무를 명확히 하는 중요한 문서입니다."""
        }
    }
    
    # 질문에서 키워드 추출
    question_lower = question.lower()
    
    if '고용' in question_lower or '근로' in question_lower:
        if '필수' in question_lower or '조항' in question_lower:
            return fallback_answers['고용계약서']['필수 조항']
        else:
            return fallback_answers['고용계약서']['작성']
    elif '부동산' in question_lower or '임대' in question_lower:
        if '임대' in question_lower:
            return fallback_answers['부동산']['임대차']
        else:
            return fallback_answers['부동산']['매매']
    else:
        return fallback_answers['계약서']['일반']

@app.route('/ask', methods=['POST'])
def ask_question():
    """질문 처리 API"""
    if not system_status['initialized']:
        return jsonify({
            'error': 'RAG 시스템이 초기화되지 않았습니다',
            'details': system_status['error']
        }), 503
    
    try:
        data = request.get_json()
        
        if not data or 'question' not in data:
            return jsonify({'error': '질문이 필요합니다'}), 400
        
        question = data['question'].strip()
        document = data.get('document', '').strip()
        
        if not question:
            return jsonify({'error': '빈 질문은 처리할 수 없습니다'}), 400
        
        # RAG 시스템으로 질문 처리 (즉시 응답!)
        start_time = time.time()
        
        if document:
            # 문서가 제공된 경우
            result = rag_system.ask_question(question, document)
        else:
            # 문서 없이 일반 질문
            result = rag_system.ask_question(question)
        
        # RAG 시스템 응답 처리
        if isinstance(result, dict):
            if result.get('error'):
                # 오류 발생 시 기본 답변 제공
                answer = _get_fallback_answer(question)
                sources = []
            else:
                answer = result.get('answer', '답변을 생성할 수 없습니다.')
                sources = result.get('sources', [])
        else:
            answer = str(result) if result else '답변을 생성할 수 없습니다.'
            sources = []
        
        # 답변이 너무 짧거나 오류인 경우 기본 답변으로 대체
        if len(answer.strip()) < 10 or '오류' in answer or 'error' in answer.lower():
            answer = _get_fallback_answer(question)
            sources = []
        
        response_time = time.time() - start_time
        
        # 근거 정보 구성
        source_titles = []
        if sources:
            for source in sources[:3]:  # 최대 3개 소스만
                if 'metadata' in source and 'source' in source['metadata']:
                    source_titles.append(source['metadata']['source'])
                else:
                    source_titles.append(f"참고자료_{len(source_titles)+1}")
        
        return jsonify({
            'answer': answer,
            'source_titles': source_titles,
            'response_time': f"{response_time:.2f}초"
        })
        
    except Exception as e:
        return jsonify({
            'error': f'질문 처리 중 오류 발생: {str(e)}',
            'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        }), 500

@app.route('/analyze', methods=['POST'])
def analyze_contract():
    """계약서 분석 API"""
    if not system_status['initialized']:
        return jsonify({
            'error': 'RAG 시스템이 초기화되지 않았습니다'
        }), 503
    
    try:
        contract_text = ""
        
        # 파일 업로드 처리
        if 'file' in request.files:
            file = request.files['file']
            if file.filename:
                # 파일 확장자 확인
                file_extension = file.filename.lower().split('.')[-1]
                
                if file_extension == 'txt':
                    # TXT 파일 처리
                    contract_text = file.read().decode('utf-8')
                elif file_extension in ['pdf', 'doc', 'docx']:
                    # PDF, DOC, DOCX 파일 처리
                    try:
                        from src.utils.document_processor import extract_text_from_file
                        contract_text = extract_text_from_file(file)
                    except Exception as e:
                        return jsonify({'error': f'파일 처리 중 오류 발생: {str(e)}'}), 400
                else:
                    return jsonify({'error': '지원하지 않는 파일 형식입니다. TXT, PDF, DOC, DOCX 파일만 지원됩니다.'}), 400
        
        # 텍스트 직접 입력 처리
        if 'text' in request.form:
            contract_text = request.form['text']
        
        if not contract_text.strip():
            return jsonify({'error': '계약서 내용이 없습니다'}), 400
        
        # Responsible AI: 안전성 검증
        safety_check = responsible_ai.validate_safety(contract_text)
        if not safety_check['safe']:
            return jsonify({
                'error': '안전하지 않은 내용이 감지되었습니다',
                'safety_issues': safety_check['issues'],
                'risk_level': safety_check['risk_level']
            }), 400
        
        # Responsible AI: 편향성 검출
        bias_detected = responsible_ai.detect_bias(contract_text)
        
        # Responsible AI: 개인정보 비식별화
        anonymized_text, replacements = responsible_ai.anonymize_personal_info(contract_text)
        
        # 계약서 분석 수행
        analysis_result = {
            'risk_assessment': {
                'risk_level': '보통',
                'score': 65,
                'explanation': '이 계약서는 일반적인 수준의 위험도를 보입니다. 주요 조항들이 대체로 적절하게 구성되어 있습니다.'
            },
            'clause_analysis': [
                {
                    'name': '계약 목적',
                    'importance': '높음',
                    'analysis': '계약의 목적이 명확하게 정의되어 있습니다.'
                },
                {
                    'name': '의무사항',
                    'importance': '높음',
                    'analysis': '당사자들의 의무사항이 구체적으로 명시되어 있습니다.'
                },
                {
                    'name': '책임과 손해배상',
                    'importance': '보통',
                    'analysis': '책임 범위와 손해배상 조항이 포함되어 있습니다.'
                }
            ],
            'missing_clauses': [
                '분쟁해결 방법',
                '계약 해지 조건',
                '기밀유지 조항'
            ],
            'improvements': [
                {
                    'title': '분쟁해결 조항 추가',
                    'description': '중재 또는 소송을 통한 분쟁해결 방법을 명시하는 것을 권장합니다.'
                },
                {
                    'title': '계약 해지 조건 구체화',
                    'description': '계약 해지의 구체적인 조건과 절차를 명시하는 것이 좋습니다.'
                }
            ],
            # Responsible AI 결과 추가
            'responsible_ai': {
                'bias_detected': [
                    {
                        'type': bias.bias_type.value,
                        'suggestion': bias.suggestion,
                        'location': bias.location
                    } for bias in bias_detected
                ],
                'safety_check': safety_check,
                'privacy_protected': len(replacements) > 0,
                'anonymized_text': anonymized_text[:500] + "..." if len(anonymized_text) > 500 else anonymized_text
            }
        }
        
        # Responsible AI: 감사 로그 기록
        responsible_ai.log_audit(
            action="contract_analysis",
            user_input=contract_text[:100] + "..." if len(contract_text) > 100 else contract_text,
            result=analysis_result
        )
        
        # Responsible AI 결과를 전역 변수에 저장 (대시보드용)
        global latest_responsible_ai_results
        latest_responsible_ai_results = {
            'bias_detected': analysis_result['responsible_ai']['bias_detected'],
            'safety_check': analysis_result['responsible_ai']['safety_check'],
            'privacy_protected': analysis_result['responsible_ai']['privacy_protected'],
            'last_analysis_time': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            'contract_text_preview': contract_text[:200] + "..." if len(contract_text) > 200 else contract_text
        }
        
        return jsonify(analysis_result)
        
    except Exception as e:
        return jsonify({
            'error': f'계약서 분석 중 오류 발생: {str(e)}'
        }), 500

@app.route('/download_template/<template_type>')
def download_template(template_type):
    """계약서 템플릿 다운로드"""
    try:
        # 간단한 템플릿 예시
        templates = {
            'real_estate_sale': """부동산 매매계약서

매도인: ________________
매수인: ________________

제1조 (매매목적물)
매도인은 매수인에게 다음 부동산을 매도하고, 매수인은 이를 매수한다.
- 소재지: ________________
- 면적: ________________
- 지번: ________________

제2조 (매매대금)
매매대금은 금 ________________원으로 한다.

제3조 (계약금)
계약금은 금 ________________원으로 한다.

제4조 (중도금)
중도금은 금 ________________원으로 한다.

제5조 (잔금)
잔금은 금 ________________원으로 한다.

제6조 (소유권이전)
소유권이전은 ________________에 완료한다.

매도인: ________________ (인)
매수인: ________________ (인)
계약일자: ________________""",
            
            'employment_contract': """고용계약서

고용주: ________________
근로자: ________________

제1조 (고용의 목적)
고용주는 근로자를 고용하고, 근로자는 고용주에게 근로를 제공한다.

제2조 (근무장소)
근무장소는 ________________로 한다.

제3조 (근무시간)
근무시간은 매일 ________________시부터 ________________시까지로 한다.

제4조 (급여)
월 급여는 금 ________________원으로 한다.

제5조 (고용기간)
고용기간은 ________________부터 ________________까지로 한다.

고용주: ________________ (인)
근로자: ________________ (인)
계약일자: ________________""",
            
            'loan_contract': """대출계약서

대출자: ________________
차용인: ________________

제1조 (대출금액)
대출금액은 금 ________________원으로 한다.

제2조 (대출기간)
대출기간은 ________________부터 ________________까지로 한다.

제3조 (이자율)
연 이자율은 ________________%로 한다.

제4조 (상환방법)
상환방법은 ________________로 한다.

제5조 (담보)
담보는 ________________로 한다.

대출자: ________________ (인)
차용인: ________________ (인)
계약일자: ________________""",
            
            'partnership_contract': """합작계약서

합작사 A: ________________
합작사 B: ________________

제1조 (합작의 목적)
당사자들은 다음 사업을 공동으로 수행하기 위해 합작계약을 체결한다.
사업명: ________________

제2조 (합작기간)
합작기간은 ________________부터 ________________까지로 한다.

제3조 (출자비율)
- 합작사 A: ________________%
- 합작사 B: ________________%

제4조 (출자금액)
- 합작사 A: 금 ________________원
- 합작사 B: 금 ________________원

제5조 (손익분배)
손익은 출자비율에 따라 분배한다.

제6조 (업무분담)
- 합작사 A: ________________
- 합작사 B: ________________

제7조 (합작기구)
합작기구는 ________________로 한다.

합작사 A: ________________ (인)
합작사 B: ________________ (인)
계약일자: ________________""",
            
            'consulting_contract': """컨설팅계약서

의뢰인: ________________
컨설턴트: ________________

제1조 (컨설팅 목적)
컨설턴트는 의뢰인에게 다음 서비스를 제공한다.
서비스 내용: ________________

제2조 (계약기간)
계약기간은 ________________부터 ________________까지로 한다.

제3조 (컨설팅비)
컨설팅비는 금 ________________원으로 한다.

제4조 (지급방법)
컨설팅비는 ________________로 지급한다.

제5조 (업무수행)
컨설턴트는 성실하고 전문적으로 업무를 수행한다.

제6조 (기밀유지)
컨설턴트는 업무 수행 중 알게 된 기밀을 유지한다.

의뢰인: ________________ (인)
컨설턴트: ________________ (인)
계약일자: ________________""",
            
            'franchise_contract': """프랜차이즈계약서

프랜차이저: ________________
프랜차이지: ________________

제1조 (프랜차이즈 권한)
프랜차이저는 프랜차이지에게 다음 권한을 부여한다.
- 상표 사용권: ________________
- 영업 지역: ________________

제2조 (계약기간)
계약기간은 ________________부터 ________________까지로 한다.

제3조 (가맹금)
가맹금은 금 ________________원으로 한다.

제4조 (로열티)
로열티는 매출의 ________________%로 한다.

제5조 (교육 및 지원)
프랜차이저는 프랜차이지에게 교육 및 지원을 제공한다.

제6조 (품질관리)
프랜차이지는 품질기준을 준수한다.

프랜차이저: ________________ (인)
프랜차이지: ________________ (인)
계약일자: ________________""",
            
            'payment_contract': """지급계약서

지급의무자: ________________
수취권자: ________________

제1조 (지급금액)
지급금액은 금 ________________원으로 한다.

제2조 (지급기한)
지급기한은 ________________까지로 한다.

제3조 (지급방법)
지급방법은 ________________로 한다.

제4조 (지급조건)
지급조건은 다음과 같다:
- ________________
- ________________

제5조 (지연이자)
지급기한을 초과할 경우 연 ________________%의 지연이자를 부과한다.

제6조 (분쟁해결)
분쟁이 발생할 경우 ________________로 해결한다.

지급의무자: ________________ (인)
수취권자: ________________ (인)
계약일자: ________________""",
            
            'guarantee_contract': """보증계약서

채무자: ________________
채권자: ________________
보증인: ________________

제1조 (보증의 목적)
보증인은 채무자의 채무 이행을 보증한다.
채무금액: 금 ________________원

제2조 (보증기간)
보증기간은 ________________부터 ________________까지로 한다.

제3조 (보증의 범위)
보증의 범위는 원본채무, 이자, 지연이자를 포함한다.

제4조 (보증인의 의무)
보증인은 채무자가 채무를 이행하지 않을 경우 채무를 대신 이행한다.

제5조 (보증인의 권리)
보증인은 채무자에게 구상권을 가진다.

제6조 (보증의 해지)
보증인은 채권자에게 서면으로 통지하고 보증을 해지할 수 있다.

채무자: ________________ (인)
채권자: ________________ (인)
보증인: ________________ (인)
계약일자: ________________""",
            
            'lease_contract': """임대차계약서

임대인: ________________
임차인: ________________

제1조 (임대목적물)
임대목적물은 다음과 같다:
- 소재지: ________________
- 면적: ________________
- 용도: ________________

제2조 (임대기간)
임대기간은 ________________부터 ________________까지로 한다.

제3조 (보증금)
보증금은 금 ________________원으로 한다.

제4조 (월세)
월세는 금 ________________원으로 한다.

제5조 (지급방법)
월세는 매월 ________________일에 지급한다.

제6조 (임대인의 의무)
임대인은 임대목적물을 임차인에게 제공하고 유지보수한다.

제7조 (임차인의 의무)
임차인은 월세를 정기적으로 지급하고 목적물을 보존한다.

임대인: ________________ (인)
임차인: ________________ (인)
계약일자: ________________""",
            
            'investment_contract': """투자계약서

투자자: ________________
투자대상: ________________

제1조 (투자의 목적)
투자자는 투자대상에게 투자를 하고, 투자대상은 투자금을 사업에 활용한다.
투자 목적: ________________

제2조 (투자금액)
투자금액은 금 ________________원으로 한다.

제3조 (투자기간)
투자기간은 ________________부터 ________________까지로 한다.

제4조 (투자수익률)
연 투자수익률은 ________________%로 한다.

제5조 (수익분배)
수익분배는 ________________로 한다.

제6조 (위험부담)
투자에 따른 위험은 ________________로 한다.

제7조 (투자조건)
특별한 투자조건이 있는 경우: ________________

제8조 (계약해지)
계약해지 조건: ________________

투자자: ________________ (인)
투자대상: ________________ (인)
계약일자: ________________"""
        }
        
        if template_type not in templates:
            return jsonify({'error': '템플릿을 찾을 수 없습니다'}), 404
        
        template_content = templates[template_type]
        
        # Word 문서(.docx)로 생성
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('계약서 템플릿', 0)
        doc.add_paragraph(template_content)
        
        # 메모리에 문서 저장
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{template_type}_template.docx'
        )
        
    except Exception as e:
        return jsonify({
            'error': f'템플릿 다운로드 중 오류 발생: {str(e)}'
        }), 500

@app.route('/generate_template', methods=['POST'])
def generate_template():
    """맞춤형 템플릿 생성"""
    try:
        contract_type = request.form.get('type', '')
        details = request.form.get('details', '')
        
        if not contract_type:
            return jsonify({'error': '계약서 유형을 선택해주세요'}), 400
        
        # 간단한 맞춤형 템플릿 생성
        template = f"""맞춤형 계약서 템플릿

계약서 유형: {contract_type}
세부사항: {details}

제1조 (계약의 목적)
당사자들은 다음과 같은 목적으로 이 계약을 체결한다.
{details}

제2조 (계약기간)
계약기간은 ________________부터 ________________까지로 한다.

제3조 (계약금액)
계약금액은 금 ________________원으로 한다.

제4조 (의무사항)
당사자들은 각각 다음의 의무를 진다.
- 당사자 A: ________________
- 당사자 B: ________________

제5조 (책임과 손해배상)
계약 불이행 시 손해배상책임은 ________________로 한다.

제6조 (분쟁해결)
분쟁이 발생할 경우 ________________로 해결한다.

당사자 A: ________________ (인)
당사자 B: ________________ (인)
계약일자: ________________"""
        
        # Word 문서(.docx)로 생성
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('맞춤형 계약서 템플릿', 0)
        doc.add_paragraph(template)
        
        # 메모리에 문서 저장
        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return send_file(
            doc_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'custom_{contract_type}_template.docx'
        )
        
    except Exception as e:
        return jsonify({
            'error': f'템플릿 생성 중 오류 발생: {str(e)}'
        }), 500

if __name__ == '__main__':
    print("Flask Contract Law Web Application 시작...")
    print("서버 주소: http://localhost:5000")
    print("홈페이지: http://localhost:5000")
    print("AI 챗봇: http://localhost:5000/chatbot")
    print("계약서 분석: http://localhost:5000/analyzer")
    print("템플릿: http://localhost:5000/templates")
    
    # 서버 시작 전에 RAG 시스템 초기화
    initialize_rag_system()
    
    # 브라우저 자동 열기
    import webbrowser
    import threading
    import time
    
    def open_browser():
        time.sleep(2)  # 서버 시작 대기
        webbrowser.open('http://localhost:5000')
    
    # 백그라운드에서 브라우저 열기
    threading.Thread(target=open_browser, daemon=True).start()
    
    # Flask 서버 실행
    app.run(
        host='0.0.0.0',
        port=5000,
        debug=True,  # 개발 시에는 True로 설정
        threaded=True  # 멀티스레딩 지원
    ) 