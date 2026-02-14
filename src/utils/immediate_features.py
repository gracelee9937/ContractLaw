"""
Immediate Features
RAG 시스템이나 벡터 스토어 없이도 즉시 사용 가능한 기능들
"""

import re
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches
import io

class ImmediateContractAnalyzer:
    """즉시 사용 가능한 계약서 분석 기능"""
    
    def __init__(self):
        # 위험도 분석을 위한 패턴들
        self.risk_patterns = {
            'high_risk': {
                '일방적 해지': [r'갑은.*해지.*수.*있다', r'해지.*권한.*갑'],
                '과도한 손해배상': [r'손해배상.*\d+배', r'위약금.*\d+%'],
                '불리한 관할법원': [r'관할.*법원.*갑.*소재지', r'소송.*갑.*관할'],
                '무제한 책임': [r'무제한.*책임', r'전액.*배상'],
                '불공정 계약조건': [r'갑.*일방적', r'을.*의무.*갑.*권리']
            },
            'medium_risk': {
                '모호한 계약기간': [r'계약기간.*협의', r'기간.*정하지.*않'],
                '불명확한 업무범위': [r'업무.*포함.*기타', r'범위.*협의.*결정'],
                '지연 손해조항': [r'지연.*손해', r'연체.*이자', r'지연배상금'],
                '비밀유지 의무': [r'비밀.*유지', r'기밀.*누설.*금지']
            },
            'low_risk': {
                '일반적 조항': [r'선량한.*관리자', r'신의성실.*원칙'],
                '표준 해지조항': [r'30일.*전.*통지', r'해지.*통지'],
                '일반 책임조항': [r'고의.*중과실', r'일반적.*손해']
            }
        }
        
        # 필수 조항 체크리스트
        self.essential_clauses = {
            '계약 당사자': [r'갑.*\(.*\)', r'을.*\(.*\)', r'당사자'],
            '계약 목적': [r'목적', r'본.*계약', r'계약.*내용'],
            '계약 기간': [r'계약.*기간', r'유효.*기간', r'기간.*\d+'],
            '대금 지급': [r'대금', r'금액', r'지급', r'수수료'],
            '의무와 책임': [r'의무', r'책임', r'준수'],
            '해지 조건': [r'해지', r'종료', r'취소'],
            '분쟁 해결': [r'분쟁', r'중재', r'관할', r'소송']
        }

    def analyze_contract_risk(self, contract_text: str) -> Dict[str, Any]:
        """계약서 위험도 즉시 분석"""
        if not contract_text:
            return {"error": "계약서 텍스트가 없습니다."}
        
        risk_score = 0
        high_risk_issues = []
        medium_risk_issues = []
        low_risk_issues = []
        missing_clauses = []
        
        # 고위험 패턴 검사
        for issue, patterns in self.risk_patterns['high_risk'].items():
            for pattern in patterns:
                if re.search(pattern, contract_text, re.IGNORECASE):
                    high_risk_issues.append(issue)
                    risk_score += 30
                    break
        
        # 중위험 패턴 검사
        for issue, patterns in self.risk_patterns['medium_risk'].items():
            for pattern in patterns:
                if re.search(pattern, contract_text, re.IGNORECASE):
                    medium_risk_issues.append(issue)
                    risk_score += 15
                    break
        
        # 저위험 패턴 검사 (긍정적 요소)
        for issue, patterns in self.risk_patterns['low_risk'].items():
            for pattern in patterns:
                if re.search(pattern, contract_text, re.IGNORECASE):
                    low_risk_issues.append(issue)
                    risk_score -= 5  # 위험도 감소
                    break
        
        # 필수 조항 확인
        for clause, patterns in self.essential_clauses.items():
            found = False
            for pattern in patterns:
                if re.search(pattern, contract_text, re.IGNORECASE):
                    found = True
                    break
            if not found:
                missing_clauses.append(clause)
                risk_score += 10
        
        # 위험도 등급 결정
        if risk_score >= 80:
            risk_level = "매우 높음"
            risk_color = "🔴"
        elif risk_score >= 50:
            risk_level = "높음"
            risk_color = "🟠"
        elif risk_score >= 25:
            risk_level = "보통"
            risk_color = "yellow"
        else:
            risk_level = "낮음"
            risk_color = "🟢"
        
        return {
            'risk_score': min(risk_score, 100),
            'risk_level': risk_level,
            'risk_color': risk_color,
            'high_risk_issues': high_risk_issues,
            'medium_risk_issues': medium_risk_issues,
            'low_risk_issues': low_risk_issues,
            'missing_clauses': missing_clauses,
            'recommendations': self._generate_recommendations(high_risk_issues, missing_clauses)
        }
    
    def _generate_recommendations(self, high_risks: List[str], missing_clauses: List[str]) -> List[str]:
        """개선 권고사항 생성"""
        recommendations = []
        
        if '일방적 해지' in high_risks:
            recommendations.append(" 해지 조건을 상호 대등하게 수정하세요")
        
        if '과도한 손해배상' in high_risks:
            recommendations.append(" 손해배상 범위를 합리적 수준으로 조정하세요")
        
        if '계약 당사자' in missing_clauses:
            recommendations.append(" 계약 당사자 정보를 명확히 기재하세요")
        
        if '계약 기간' in missing_clauses:
            recommendations.append(" 계약 기간을 구체적으로 명시하세요")
        
        if '분쟁 해결' in missing_clauses:
            recommendations.append(" 분쟁 해결 방법을 포함하세요")
        
        if not recommendations:
            recommendations.append(" 전반적으로 양호한 계약서입니다")
        
        return recommendations

class ImmediateTemplateGenerator:
    """즉시 사용 가능한 계약서 템플릿 생성기"""
    
    def __init__(self):
        self.templates = {
            '부동산 매매계약서': {
                'title': '부동산 매매계약서',
                'sections': [
                    ('제1조 (계약목적물)', '본 계약의 목적물은 다음과 같다.\n위치: [주소 기재]\n면적: [면적 기재]\n기타: [기타 사항]'),
                    ('제2조 (매매대금)', '매매대금은 금 [금액] 원으로 한다.\n계약금: [계약금]\n중도금: [중도금]\n잔금: [잔금]'),
                    ('제3조 (계약이행)', '잔금지급일: [날짜]\n소유권이전: 잔금 지급과 동시에 이행'),
                    ('제4조 (계약해제)', '당사자 일방이 본 계약을 위반한 경우 상대방은 계약을 해제할 수 있다.'),
                    ('제5조 (분쟁해결)', '본 계약으로 인한 분쟁은 [관할법원]의 관할로 한다.')
                ]
            },
            '고용계약서': {
                'title': '고용계약서',
                'sections': [
                    ('제1조 (근무조건)', '근무지: [근무지 기재]\n직위: [직위 기재]\n근무시간: [근무시간]'),
                    ('제2조 (급여)', '기본급: 월 [금액] 원\n상여금: [상여금 조건]\n지급일: 매월 [날짜]'),
                    ('제3조 (근무의무)', '직원은 성실히 업무를 수행하여야 한다.'),
                    ('제4조 (계약기간)', '계약기간: [시작일]부터 [종료일]까지'),
                    ('제5조 (계약해지)', '30일 전 서면 통지로 계약을 해지할 수 있다.')
                ]
            },
            '용역계약서': {
                'title': '용역계약서',
                'sections': [
                    ('제1조 (용역의 내용)', '용역의 종류: [용역 내용]\n수행기간: [기간]\n수행장소: [장소]'),
                    ('제2조 (용역대금)', '총 용역대금: 금 [금액] 원\n지급방법: [지급 조건]'),
                    ('제3조 (의무사항)', '수급인은 선량한 관리자의 주의로 용역을 수행한다.'),
                    ('제4조 (완성 및 인도)', '용역 완성일: [날짜]\n검수 및 인수: [조건]'),
                    ('제5조 (손해배상)', '고의 또는 중과실로 인한 손해에 대해 배상책임을 진다.')
                ]
            },
            '매매계약서': {
                'title': '매매계약서',
                'sections': [
                    ('제1조 (매매목적물)', '품명: [품명]\n규격: [규격]\n수량: [수량]\n단가: [단가]'),
                    ('제2조 (대금)', '총 대금: 금 [총액] 원\n지급방법: [지급 조건]'),
                    ('제3조 (인도)', '인도장소: [장소]\n인도기한: [날짜]\n위험부담: 인도시 이전'),
                    ('제4조 (검사)', '매수인은 목적물 수령 후 7일 이내 검사하여야 한다.'),
                    ('제5조 (하자담보)', '매도인은 인도 후 6개월간 하자담보책임을 진다.')
                ]
            }
        }
    
    def generate_template(self, template_type: str, custom_data: Dict[str, str] = None) -> bytes:
        """계약서 템플릿 생성"""
        if template_type not in self.templates:
            raise ValueError(f"지원하지 않는 템플릿 유형: {template_type}")
        
        template = self.templates[template_type]
        
        # 워드 문서 생성
        doc = Document()
        
        # 제목
        title = doc.add_heading(template['title'], 0)
        title.alignment = 1  # 중앙 정렬
        
        # 현재 날짜
        doc.add_paragraph(f"작성일: {datetime.now().strftime('%Y년 %m월 %d일')}")
        doc.add_paragraph()
        
        # 당사자
        doc.add_paragraph("갑: [회사/개인명] (이하 '갑'이라 한다)")
        doc.add_paragraph("을: [회사/개인명] (이하 '을'이라 한다)")
        doc.add_paragraph()
        
        doc.add_paragraph("갑과 을은 다음과 같이 계약을 체결한다.")
        doc.add_paragraph()
        
        # 계약 조항들
        for section_title, section_content in template['sections']:
            doc.add_heading(section_title, level=1)
            
            # 커스텀 데이터가 있으면 치환
            if custom_data:
                for key, value in custom_data.items():
                    section_content = section_content.replace(f'[{key}]', value)
            
            doc.add_paragraph(section_content)
            doc.add_paragraph()
        
        # 서명란
        doc.add_paragraph()
        doc.add_paragraph("갑: ___________________ (서명/인)")
        doc.add_paragraph("을: ___________________ (서명/인)")
        
        # 바이트로 변환
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def get_available_templates(self) -> List[str]:
        """사용 가능한 템플릿 목록 반환"""
        return list(self.templates.keys())

class ImmediateClauseExtractor:
    """즉시 사용 가능한 조항 추출기"""
    
    def __init__(self):
        self.clause_patterns = {
            '계약 당사자': [
                r'갑.*?(?=\n|\r|\.|을)',
                r'을.*?(?=\n|\r|\.|갑)',
                r'당사자.*?(?=\n|\r|\.)'
            ],
            '계약 기간': [
                r'계약.*?기간.*?(?=\n|\r|\.)',
                r'\d{4}년.*?\d{1,2}월.*?\d{1,2}일.*?부터.*?\d{4}년.*?\d{1,2}월.*?\d{1,2}일.*?까지',
                r'기간.*?\d+.*?(?:개월|년|일)'
            ],
            '대금 및 지급조건': [
                r'대금.*?금.*?\d+.*?원',
                r'지급.*?방법.*?(?=\n|\r|\.)',
                r'매매.*?대금.*?(?=\n|\r|\.)'
            ],
            '해지 조건': [
                r'해지.*?(?=\n|\r|제\d+조)',
                r'계약.*?종료.*?(?=\n|\r|\.)',
                r'취소.*?(?=\n|\r|\.)'
            ],
            '분쟁 해결': [
                r'분쟁.*?(?=\n|\r|제\d+조)',
                r'관할.*?법원.*?(?=\n|\r|\.)',
                r'소송.*?(?=\n|\r|\.)',
                r'중재.*?(?=\n|\r|\.)'
            ]
        }
    
    def extract_clauses(self, contract_text: str) -> Dict[str, List[str]]:
        """계약서에서 주요 조항 추출"""
        extracted_clauses = {}
        
        for clause_type, patterns in self.clause_patterns.items():
            matches = []
            for pattern in patterns:
                found_matches = re.findall(pattern, contract_text, re.IGNORECASE | re.MULTILINE)
                matches.extend(found_matches)
            
            # 중복 제거 및 정리
            unique_matches = list(set(matches))
            cleaned_matches = [match.strip() for match in unique_matches if match.strip()]
            
            extracted_clauses[clause_type] = cleaned_matches[:3]  # 최대 3개까지
        
        return extracted_clauses

# 전역 인스턴스들
immediate_analyzer = ImmediateContractAnalyzer()
immediate_template_generator = ImmediateTemplateGenerator()
immediate_clause_extractor = ImmediateClauseExtractor() 