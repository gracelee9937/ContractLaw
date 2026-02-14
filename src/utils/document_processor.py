"""
Document Processing Utilities for ContractLawProject
Handles PDF processing, text extraction, and contract parsing
"""
import os
import re
import zipfile
import tempfile
from typing import Dict, List, Tuple, Optional
import fitz  # PyMuPDF
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class DocumentProcessor:
    """Main document processing class"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF file"""
        try:
            if hasattr(pdf_file, 'read'):
                pdf_data = pdf_file.read()
                pdf_file.seek(0)  # Reset file pointer
            else:
                with open(pdf_file, 'rb') as f:
                    pdf_data = f.read()
            
            doc = fitz.open(stream=pdf_data, filetype="pdf")
            text = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text += page.get_text()
            
            doc.close()
            return text.strip()
        
        except Exception as e:
            print(f"ERROR: PDF 처리 중 오류가 발생했습니다: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file"""
        try:
            if hasattr(docx_file, 'read'):
                doc = Document(docx_file)
            else:
                doc = Document(docx_file)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text.strip()
        
        except Exception as e:
            print(f"ERROR: DOCX 처리 중 오류가 발생했습니다: {str(e)}")
            return ""


def extract_text_from_file(uploaded_file) -> str:
    """Extract text from various file formats - 독립적인 함수"""
    if uploaded_file is None:
        return ""
    
    file_extension = os.path.splitext(uploaded_file.filename)[1].lower()
    
    doc_processor = DocumentProcessor()
    
    if file_extension == '.pdf':
        return doc_processor.extract_text_from_pdf(uploaded_file)
    elif file_extension == '.docx':
        return doc_processor.extract_text_from_docx(uploaded_file)
    elif file_extension == '.txt':
        return str(uploaded_file.read(), "utf-8")
    else:
        print(f"ERROR: 지원되지 않는 파일 형식입니다: {file_extension}")
        return ""


    def extract_text_from_file(self, uploaded_file) -> str:
        """Extract text from various file formats"""
        if uploaded_file is None:
            return ""
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension == '.pdf':
            return self.extract_text_from_pdf(uploaded_file)
        elif file_extension == '.docx':
            return self.extract_text_from_docx(uploaded_file)
        elif file_extension == '.txt':
            return str(uploaded_file.read(), "utf-8")
        else:
            print(f"ERROR: 지원되지 않는 파일 형식입니다: {file_extension}")
            return ""

class ContractAnalyzer:
    """Contract analysis and clause extraction"""
    
    def __init__(self):
        # Korean contract keywords for clause identification
        self.clause_patterns = {
            "payment_terms": [
                r"지급\s*조건", r"대금\s*지급", r"결제\s*방법", r"지불\s*조건", 
                r"대금\s*결제", r"지급\s*기한", r"납부\s*기한"
            ],
            "termination_clauses": [
                r"해지\s*조건", r"계약\s*해지", r"종료\s*조건", r"계약\s*종료",
                r"해약\s*조건", r"중도\s*해지"
            ],
            "liability_limitations": [
                r"책임\s*제한", r"면책\s*조항", r"손해\s*배상", r"책임\s*한계",
                r"면책\s*사항", r"배상\s*책임"
            ],
            "intellectual_property": [
                r"지적\s*재산권", r"저작권", r"특허권", r"상표권", r"영업비밀",
                r"기술\s*자료", r"노하우"
            ],
            "confidentiality": [
                r"비밀\s*유지", r"기밀\s*유지", r"비밀\s*보장", r"기밀\s*보호",
                r"비밀\s*누설\s*금지", r"기밀\s*정보"
            ],
            "dispute_resolution": [
                r"분쟁\s*해결", r"중재\s*조항", r"관할\s*법원", r"준거법",
                r"조정\s*절차", r"중재\s*기관"
            ],
            "force_majeure": [
                r"불가항력", r"천재지변", r"불가피한\s*사유", r"부득이한\s*사유"
            ],
            "compliance_requirements": [
                r"법령\s*준수", r"규정\s*준수", r"컴플라이언스", r"준법",
                r"관련\s*법규", r"적용\s*법률"
            ]
        }
    
    def extract_clauses(self, contract_text: str) -> Dict[str, List[str]]:
        """Extract specific clauses from contract text"""
        found_clauses = {}
        
        for clause_type, patterns in self.clause_patterns.items():
            found_clauses[clause_type] = []
            
            for pattern in patterns:
                matches = re.finditer(pattern, contract_text, re.IGNORECASE)
                for match in matches:
                    # Extract surrounding context (±200 characters)
                    start = max(0, match.start() - 200)
                    end = min(len(contract_text), match.end() + 200)
                    context = contract_text[start:end].strip()
                    
                    if context not in found_clauses[clause_type]:
                        found_clauses[clause_type].append(context)
        
        return found_clauses
    
    def calculate_risk_score(self, clauses: Dict[str, List[str]]) -> Dict[str, float]:
        """Calculate risk scores based on clause presence and content"""
        risk_scores = {}
        
        # Base weights for different clause types
        clause_weights = {
            "payment_terms": 0.9,
            "termination_clauses": 0.8,
            "liability_limitations": 0.7,
            "intellectual_property": 0.6,
            "confidentiality": 0.6,
            "dispute_resolution": 0.8,
            "force_majeure": 0.5,
            "compliance_requirements": 0.7
        }
        
        total_weighted_score = 0
        total_weight = 0
        
        for clause_type, weight in clause_weights.items():
            if clause_type in clauses and clauses[clause_type]:
                # Score based on presence and detail level
                num_clauses = len(clauses[clause_type])
                avg_length = sum(len(clause) for clause in clauses[clause_type]) / num_clauses
                
                # Scoring logic: more detailed clauses = lower risk
                if avg_length > 300:
                    clause_score = 0.2  # Low risk - detailed clauses
                elif avg_length > 150:
                    clause_score = 0.5  # Medium risk
                else:
                    clause_score = 0.7  # Higher risk - brief clauses
            else:
                clause_score = 1.0  # High risk - missing clause
            
            risk_scores[clause_type] = clause_score
            total_weighted_score += clause_score * weight
            total_weight += weight
        
        # Overall risk score (0-1, where 1 is highest risk)
        risk_scores['overall'] = total_weighted_score / total_weight if total_weight > 0 else 1.0
        
        return risk_scores
    
    def get_missing_clauses(self, clauses: Dict[str, List[str]]) -> List[str]:
        """Identify missing critical clauses"""
        missing = []
        
        critical_clauses = {
            "payment_terms": "결제 조건",
            "termination_clauses": "해지 조건", 
            "liability_limitations": "책임 제한",
            "dispute_resolution": "분쟁 해결"
        }
        
        for clause_type, korean_name in critical_clauses.items():
            if clause_type not in clauses or not clauses[clause_type]:
                missing.append(korean_name)
        
        return missing
    
    def summarize_contract(self, contract_text: str) -> Dict[str, str]:
        """Generate contract summary"""
        # Basic contract information extraction
        summary = {
            "length": len(contract_text),
            "word_count": len(contract_text.split()),
            "paragraph_count": len([p for p in contract_text.split('\n') if p.strip()]),
        }
        
        # Extract key entities (simplified pattern matching)
        parties_pattern = r'(갑|을|병|정)\s*(?:은|는|이|가|을|를)?'
        parties = list(set(re.findall(parties_pattern, contract_text)))
        summary["parties"] = parties
        
        # Extract dates
        date_patterns = [
            r'\d{4}년\s*\d{1,2}월\s*\d{1,2}일',
            r'\d{4}-\d{1,2}-\d{1,2}',
            r'\d{4}\.\d{1,2}\.\d{1,2}'
        ]
        
        dates = []
        for pattern in date_patterns:
            dates.extend(re.findall(pattern, contract_text))
        summary["dates"] = list(set(dates))
        
        return summary

class TemplateGenerator:
    """Generate contract templates"""
    
    def __init__(self):
        self.template_styles = {
            "title": {"size": 16, "bold": True, "align": WD_ALIGN_PARAGRAPH.CENTER},
            "heading": {"size": 14, "bold": True},
            "body": {"size": 12, "bold": False},
            "clause": {"size": 11, "bold": False}
        }
    
    def create_basic_contract_template(self, contract_type: str, parties: List[str]) -> Document:
        """Create a basic contract template"""
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{contract_type} 계약서', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Parties section
        doc.add_heading('제1조 (당사자)', level=2)
        for i, party in enumerate(parties, 1):
            p = doc.add_paragraph(f'{party}: ')
            p.add_run('(이하 "○○○"라 한다)')
        
        # Standard clauses
        standard_clauses = [
            ('제2조 (목적)', '본 계약은 ○○○에 관한 사항을 정함을 목적으로 한다.'),
            ('제3조 (계약기간)', '본 계약의 유효기간은 ○○○년 ○월 ○일부터 ○○○년 ○월 ○일까지로 한다.'),
            ('제4조 (대금 및 지급방법)', '○○○에 대한 대금은 ○○○원으로 하며, ○○○ 방법으로 지급한다.'),
            ('제5조 (의무사항)', '각 당사자의 주요 의무사항은 다음과 같다.'),
            ('제6조 (해지조건)', '다음 각 호의 사유가 발생할 경우 본 계약을 해지할 수 있다.'),
            ('제7조 (손해배상)', '계약 위반으로 인한 손해배상에 관한 사항을 정한다.'),
            ('제8조 (분쟁해결)', '본 계약과 관련된 분쟁은 ○○ 관할법원에서 해결한다.'),
            ('제9조 (기타)', '본 계약에 명시되지 않은 사항은 관련 법령을 따른다.')
        ]
        
        for heading, content in standard_clauses:
            doc.add_heading(heading, level=2)
            doc.add_paragraph(content)
        
        # Signature section
        doc.add_page_break()
        doc.add_heading('서명', level=2)
        
        for party in parties:
            doc.add_paragraph(f'{party}:')
            doc.add_paragraph('성명: _________________ (인)')
            doc.add_paragraph('주소: _________________')
            doc.add_paragraph('')
        
        return doc
    
    def save_template_as_docx(self, doc: Document, filename: str) -> str:
        """Save template as DOCX file"""
        import re
        from datetime import datetime
        
        # Create directory
        os.makedirs("templates/generated", exist_ok=True)
        
        # Clean filename: remove special characters and spaces
        clean_filename = re.sub(r'[^\w\-_.]', '_', filename)
        clean_filename = re.sub(r'_+', '_', clean_filename)  # Replace multiple underscores with single
        
        # Add timestamp to ensure uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        safe_filename = f"{clean_filename}_{timestamp}.docx"
        
        filepath = os.path.join("templates", "generated", safe_filename)
        
        try:
            doc.save(filepath)
            print(f"템플릿 저장 완료: {filepath}")
            return filepath
        except Exception as e:
            print(f"템플릿 저장 실패: {str(e)}")
            # Fallback: save to temp directory
            import tempfile
            temp_path = os.path.join(tempfile.gettempdir(), safe_filename)
            doc.save(temp_path)
            print(f"임시 디렉토리에 저장됨: {temp_path}")
            return temp_path

# Utility functions
def extract_zip_content(zip_path: str) -> List[str]:
    """Extract text content from ZIP files containing legal documents"""
    extracted_texts = []
    
    try:
        with zipfile.ZipFile(zip_path, 'r') as zip_ref:
            with tempfile.TemporaryDirectory() as temp_dir:
                zip_ref.extractall(temp_dir)
                
                for root, dirs, files in os.walk(temp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        if file.lower().endswith(('.txt', '.docx', '.json')):
                            try:
                                if file.lower().endswith('.txt'):
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        extracted_texts.append(f.read())
                                elif file.lower().endswith('.docx'):
                                    doc = Document(file_path)
                                    text = '\n'.join([p.text for p in doc.paragraphs])
                                    extracted_texts.append(text)
                                elif file.lower().endswith('.json'):
                                    import json
                                    with open(file_path, 'r', encoding='utf-8') as f:
                                        json_data = json.load(f)
                                        # 한국 법률 데이터셋 구조에 맞게 텍스트 추출
                                        if isinstance(json_data, dict) and 'document' in json_data:
                                            document = json_data['document']
                                            
                                            # 제목과 기본 정보 추출
                                            text_parts = []
                                            if 'title' in document:
                                                text_parts.append(f"제목: {document['title']}")
                                            if 'purpose' in document:
                                                text_parts.append(f"목적: {document['purpose']}")
                                            
                                            # sub_documents에서 실제 계약서 내용 추출
                                            if 'sub_documents' in document and isinstance(document['sub_documents'], list):
                                                for sub_doc in document['sub_documents']:
                                                    if isinstance(sub_doc, dict):
                                                        # contents 필드에서 텍스트 추출 (리스트 형태)
                                                        if 'contents' in sub_doc and sub_doc['contents']:
                                                            contents_list = sub_doc['contents']
                                                            if isinstance(contents_list, list):
                                                                for content_item in contents_list:
                                                                    if isinstance(content_item, dict) and 'text' in content_item:
                                                                        text = content_item['text']
                                                                        if isinstance(text, str) and len(text.strip()) > 5:
                                                                            text_parts.append(text.strip())
                                                            elif isinstance(contents_list, str):
                                                                # 혹시 문자열인 경우도 처리
                                                                if len(contents_list.strip()) > 5:
                                                                    text_parts.append(contents_list.strip())
                                                        
                                                        # name 필드도 포함 (조항 제목 등)
                                                        if 'name' in sub_doc and sub_doc['name']:
                                                            name = sub_doc['name']
                                                            if isinstance(name, str) and len(name.strip()) > 3:
                                                                text_parts.append(name.strip())
                                            
                                            if text_parts:
                                                extracted_texts.append('\n'.join(text_parts))
                            except Exception as e:
                                print(f"Error processing {file}: {e}")
                                continue
    
    except Exception as e:
        print(f"Error extracting ZIP {zip_path}: {e}")
    
    return extracted_texts

def clean_korean_text(text: str) -> str:
    """Clean and normalize Korean text"""
    if not text:
        return ""
    
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text.strip())
    
    # Remove special characters but keep Korean characters and basic punctuation
    text = re.sub(r'[^\w\s가-힣.,;:!?()\-]', '', text)
    
    return text 